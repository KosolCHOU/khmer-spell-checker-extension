from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json
import re
import io
import unicodedata
from .transliterator import get_transliterator
from .pdf_repair import repair_khmer_text
from .spell_checker_advanced import check_spelling

def home(request):
    return render(request, 'landing/index.html')

def spell_checker_page(request):
    return render(request, 'landing/spell-checker.html')

@csrf_exempt
def spell_check_api(request):
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    try:
        data = json.loads(request.body)
        text = (data.get('text') or '')
        if not text:
            return JsonResponse({'error': 'No text provided'}, status=400)
        
        # Use the advanced spell checker
        result = check_spelling(text)
        return JsonResponse({'success': True, **result})
    except Exception as e:
        import traceback
        print('Spell check error:', e)
        print(traceback.format_exc())
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

def fix_visual_order(text):
    """
    Fixes common Khmer PDF extraction issues where vowels 
    appearing to the left (e, ae, ai) are extracted before the consonant.
    Moves them to the correct logical position (after the full consonant cluster).
    """
    if not text:
        return ""
        
    # Pattern: Matches Pre-base Vowels (E, AE, AI) followed by a Consonant Cluster
    # Group 1: The misplaced vowel (\u17c1-\u17c3)
    # Group 2: The Consonant + optional Subscripts (Coeng + Consonant)
    # logic: [\u1780-\u17a2] is range of consonants. \u17d2 is Coeng.
    pattern = re.compile(r'([\u17c1-\u17c3])([\u1780-\u17a2](?:\u17d2[\u1780-\u17a2])*)')
    
    # Swap: Put the consonant cluster (Group 2) before the vowel (Group 1)
    fixed_text = pattern.sub(r'\2\1', text)
    
    return fixed_text



@csrf_exempt
def extract_file_text(request):
    """Extract text from uploaded files (txt, docx, pdf)"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    
    try:
        if 'file' not in request.FILES:
            return JsonResponse({'error': 'No file uploaded'}, status=400)
        
        file = request.FILES['file']
        file_name = file.name.lower()
        
        # Extract text based on file type
        text = ''
        
        if file_name.endswith('.txt'):
            # Plain text file - try UTF-8 first, then fallback to other encodings
            file_content = file.read()
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # Try other common encodings for Khmer text
                for encoding in ['utf-16', 'iso-8859-1', 'cp1252', 'utf-8-sig']:
                    try:
                        text = file_content.decode(encoding)
                        break
                    except (UnicodeDecodeError, LookupError):
                        continue
                else:
                    # Last resort: decode with error handling
                    text = file_content.decode('utf-8', errors='replace')
        
        elif file_name.endswith('.docx'):
            # Word document - extracts text and preserves formatting
            try:
                from docx import Document
                doc = Document(io.BytesIO(file.read()))
                # Extract text from paragraphs and tables
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                # Also extract from tables if present
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text)
                text = '\n'.join(text_parts)
            except ImportError:
                return JsonResponse({
                    'error': 'python-docx library not installed. Please install it to process .docx files.'
                }, status=500)
        
        elif file_name.endswith('.pdf'):
            # Method: Google Vision API for High-Accuracy Khmer OCR
            # Based on cloud_ocr.py
            try:
                import os
                from google.cloud import vision
                from pdf2image import convert_from_bytes

                # Set your Google Credentials
                # Note: Using absolute path as requested/found
                os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "/home/kosol/khmer_spelling/gen-lang-client-0935690938-987476b7034c.json"

                def detect_khmer_text(image_content):
                    client = vision.ImageAnnotatorClient()
                    image = vision.Image(content=image_content)
                    # We use document_text_detection for dense text/PDFs
                    response = client.document_text_detection(image=image)
                    return response.full_text_annotation.text

                def normalize_text(s: str) -> str:
                    # Normalize to NFC helps Khmer combining marks
                    s = unicodedata.normalize('NFC', s)
                    # Strip common control characters except newlines and tabs
                    s = ''.join(c for c in s if (c >= ' ' or c in ('\n', '\t')))
                    # Collapse excessive blank lines
                    lines = [line.rstrip() for line in s.splitlines()]
                    return '\n'.join(lines).strip()

                file_content = file.read()

                # Convert PDF to Images
                # Using 300 DPI as per cloud_ocr.py
                images = convert_from_bytes(file_content, dpi=300)
                
                extracted_parts = []
                for i, img in enumerate(images):
                    # Convert PIL image to bytes for Google API
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='JPEG')
                    content = img_byte_arr.getvalue()
                    
                    try:
                        text = detect_khmer_text(content)
                        if text:
                            extracted_parts.append(text)
                    except Exception as page_err:
                        print(f"OCR Error on page {i+1}: {page_err}")
                        # Continue to next page if one fails

                extracted = "\n\n".join(extracted_parts)

                # --- PROCESSING & REPAIR ---
                
                # 1. Normalize
                text = normalize_text(extracted)

                # 2. Repair common legacy mapping issues
                text = repair_khmer_text(text)
                
                # 3. Fix Visual Order
                text = fix_visual_order(text)
                
                if not text.strip():
                    return JsonResponse({'error': 'No text could be extracted from the PDF'}, status=400)

                return JsonResponse({
                    'success': True, 
                    'text': text, 
                    'filename': file.name,
                    'diagnostics': {'method': 'GoogleVisionOCR'}
                })

            except Exception as err:
                print(f'PDF extraction error: {err}')
                import traceback
                traceback.print_exc()
                return JsonResponse({'error': f'Error extracting PDF: {str(err)}'}, status=500)
        
        elif file_name.endswith('.doc'):
            # Old .doc format (not supported without additional libraries)
            return JsonResponse({
                'error': 'Legacy .doc format is not supported. Please convert to .docx or .pdf first.'
            }, status=400)
        
        else:
            return JsonResponse({'error': 'Unsupported file type'}, status=400)
        
        if not text.strip():
            return JsonResponse({'error': 'No text could be extracted from the file'}, status=400)
        
        return JsonResponse({
            'success': True,
            'text': text.strip(),
            'filename': file.name
        })
        
    except Exception as e:
        import traceback
        print('File extraction error:', e)
        print(traceback.format_exc())
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
def transliterate(request):
    """Handle SingKhmer to Khmer transliteration requests"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    
    try:
        data = json.loads(request.body)
        singkhmer = data.get('text', '').strip()
        
        if not singkhmer:
            return JsonResponse({'error': 'No text provided'}, status=400)
        
        # Get transliterator instance
        transliterator = get_transliterator()
        
        # Get top 3 candidates
        candidates = transliterator.translate(singkhmer, top_k=3)
        
        # Ensure we always have 3 candidates (pad with empty if needed)
        while len(candidates) < 3:
            candidates.append('')
        
        # Return with best candidate in the middle (iOS style)
        # Order: [2nd best, 1st best, 3rd best]
        response = {
            'candidates': [
                candidates[1] if len(candidates) > 1 else '',  # Left (2nd best)
                candidates[0] if len(candidates) > 0 else '',  # Middle (best)
                candidates[2] if len(candidates) > 2 else ''   # Right (3rd best)
            ],
            'success': True
        }
        
        return JsonResponse(response)
        
    except Exception as e:
        import traceback
        print(f"Transliteration error: {e}")
        print(traceback.format_exc())
        return JsonResponse({'error': str(e), 'success': False}, status=500)


